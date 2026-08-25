import base64
from io import BytesIO

from fastapi import UploadFile
from openai import AsyncOpenAI


def extract_pdf_text(data: bytes, max_pages: int = 50) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            break
        parts.append(page.extract_text() or "")
    text = "\n\n".join(parts).strip()
    if len(text) < 100 and len(reader.pages) >= 5:
        return "[это скан, OCR пока не поддерживается]"
    return text


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


async def whisper_transcribe(
    audio_bytes: bytes, filename: str, llm_client: AsyncOpenAI
) -> str:
    f = BytesIO(audio_bytes)
    f.name = filename  # SDK определяет формат по расширению
    result = await llm_client.audio.transcriptions.create(
        model="whisper-1", file=f
    )
    return result.text


async def media_to_part(media: UploadFile, llm_client: AsyncOpenAI) -> dict:
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(
            data, media.filename or "audio.ogg", llm_client
        )
        return {"type": "text", "text": f"[пользователь сказал голосом]:\n{transcript}"}

    if mime == "application/pdf":
        return {
            "type": "text",
            "text": f"[документ PDF]:\n{extract_pdf_text(data)[:30_000]}",
        }

    if mime.endswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{extract_docx_text(data)[:30_000]}",
        }

    raise ValueError(f"Unsupported media type: {mime}")
